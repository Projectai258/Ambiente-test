import streamlit as st
import openai
import os
import re
import logging
import io
import json
from dotenv import load_dotenv
from bs4 import BeautifulSoup
import markdown
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF
from pydantic import BaseModel

# Configurazione iniziale
########################################

# 1) Configurazione Streamlit (DEVE ESSERE LA PRIMA ISTRUZIONE DOPO LE IMPORT)
st.set_page_config(
    page_title="Revisione Documenti",
    layout="wide",
    initial_sidebar_state="expanded",
    page_icon="📄"
)

# 2) Carica variabili d'ambiente
load_dotenv()

class Settings(BaseModel):
    OPENROUTER_API_KEY: str

# Carica le variabili d'ambiente
settings = Settings(OPENROUTER_API_KEY=os.getenv("OPENROUTER_API_KEY"))

# 3) Configurazione logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler("app.log"), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# 4) Recupera la chiave API
API_KEY = settings.OPENROUTER_API_KEY
if not API_KEY:
    st.error("⚠️ Errore: API Key di OpenRouter non trovata! Impostala come variabile d'ambiente.")
    st.stop()

# Verifica della chiave API
try:
    client = openai.OpenAI(api_key=API_KEY, base_url="https://openrouter.ai/api/v1")
    test_response = client.chat.completions.create(
        model="google/gemini-2.0-pro-exp-02-05:free",
        messages=[{"role": "system", "content": "Test"}]
    )
    if not test_response:
        st.error("⚠️ Errore: Chiave API non valida!")
        st.stop()
except Exception as e:
    st.error(f"⚠️ Errore di connessione all'API: {e}")
    st.stop()

# Pattern critici per la revisione (aggiornati per includere anche frasi come "io e ...")
CRITICAL_PATTERNS = [
    r"\bIlias Contreas\b",
    r"\bIlias\b",
    r"\bContreas\b",
    r"\bJoey\b",
    r"\bMya\b",
    r"\bmia moglie\b",
    r"\bmia figlia\b",
    r"\bShake Your English\b",
    r"\bBarman PR\b",
    r"\bStairs Club\b",
    r"\bil mio socio\b",
    r"\bio e il mio socio\b",
    r"\bil mio corso\b",
    r"\bla mia accademia\b",
    r"\bintervista\b",
    r"Mi chiamo .*? la mia esperienza personale\.",
    r"\bflair\b",
    r"\bfiglio di papà\b",
    r"\bhappy our\b",
    r"\bio e\b\s+.+"
]
compiled_patterns = [re.compile(p, re.IGNORECASE) for p in CRITICAL_PATTERNS]

# Opzioni di tono per la riscrittura
TONE_OPTIONS = {
    "Stile originale": "Mantieni lo stesso stile e struttura del testo originale.",
    "Formale": "Riscrivi in modo formale e professionale, adatto a documenti ufficiali.",
    "Informale": "Riscrivi in modo amichevole e colloquiale, adatto a comunicazioni informali.",
    "Tecnico": "Riscrivi con linguaggio tecnico e preciso, adatto ad un manuale tecnico.",
    "Narrativo": "Riscrivi in modo descrittivo e coinvolgente stile racconto.",
    "Pubblicitario": "Riscrivi in modo persuasivo, come una pubblicità.",
    "Giornalistico": "Riscrivi in tono chiaro e informativo.",
}

# Funzioni di supporto
########################################

def ai_convert_first_singular_to_plural(text):
    if not text.strip():
        return ""
    prompt = (
        "Riscrivi il seguente testo modificando esclusivamente il modo di interloquire da prima persona singolare a prima persona plurale. "
        "Mantieni invariato il contenuto e il senso logico.\n\n"
        f"Testo originale:\n{text}"
    )
    try:
        response = client.chat.completions.create(
            model="google/gemini-2.0-pro-exp-02-05:free",
            messages=[{"role": "system", "content": prompt}],
            max_tokens=500,
            timeout=10
        )
        if response and hasattr(response, "choices") and response.choices:
            return response.choices[0].message.content.strip()
        logger.error("⚠️ Errore: Nessun testo valido restituito dall'API per la conversione in plurale.")
        return ""
    except Exception as e:
        logger.error(f"⚠️ Errore nell'elaborazione (conversione in plurale): {e}")
        return ""

def convert_plain_text_to_minimal_html(text):
    paragraphs = "".join(f"<p>{line.strip()}</p>" for line in text.splitlines() if line.strip())
    return f"""<!DOCTYPE html>
<html lang="it">
<head>
  <meta charset="utf-8">
  <title>Documento Revisionato</title>
</head>
<body>
{paragraphs}
</body>
</html>"""

def extract_context(blocks, selected_block):
    try:
        index = blocks.index(selected_block)
        prev_block = blocks[index - 1] if index > 0 else ""
        next_block = blocks[index + 1] if index < len(blocks) - 1 else ""
        return prev_block, next_block
    except ValueError:
        logger.error("Il blocco selezionato non è presente nella lista.")
        return "", ""

def ai_rewrite_text(text, prev_text, next_text, tone):
    prompt = (
        f"Contesto:\nPrecedente: {prev_text}\nTesto: {text}\nSuccessivo: {next_text}\n\n"
        f"Riscrivi il 'Testo' in tono '{tone}'. Rimuovi eventuali dettagli personali o identificabili. "
        "Rispondi con UNA sola frase, senza ulteriori commenti."
    )
    try:
        response = client.chat.completions.create(
            model="google/gemini-2.0-pro-exp-02-05:free",
            messages=[{"role": "system", "content": prompt}],
            max_tokens=50
        )
        if response and hasattr(response, "choices") and response.choices:
            return response.choices[0].message.content.strip()
        logger.error("⚠️ Errore: Nessun testo valido restituito dall'API per la riscrittura del blocco.")
        return ""
    except Exception as e:
        logger.error(f"⚠️ Errore nell'elaborazione (riscrittura del blocco): {e}")
        return ""

def ai_analyze_block(prev_text, text, next_text):
    """
    Analizza il blocco di testo, considerando il contesto (precedente e successivo),
    per valutare la presenza di informazioni sensibili, con particolare attenzione a frasi che iniziano con "io e".
    
    La risposta attesa deve essere in formato JSON esattamente come:
    {
      "classificazione": "Critico" o "Non critico",
      "motivazione": "Descrizione sintetica degli elementi problematici, se presenti"
    }
    """
    prompt = f"""Contesto:
Precedente: {prev_text}
Testo: {text}
Successivo: {next_text}

Analizza il blocco di testo e indica se è "Critico" o "Non critico" in base alla presenza di informazioni sensibili, dati personali o riferimenti problematici, con particolare attenzione a frasi che iniziano con "io e". 
Rispondi esattamente in questo formato JSON:
{{
  "classificazione": "Critico" o "Non critico",
  "motivazione": "Descrizione sintetica degli elementi problematici, se presenti."
}}
"""
    try:
        response = client.chat.completions.create(
            model="google/gemini-2.0-pro-exp-02-05:free",
            messages=[{"role": "system", "content": prompt}],
            max_tokens=150
        )
        raw_output = response.choices[0].message.content.strip() if (response and hasattr(response, "choices") and response.choices) else ""
        logger.info(f"Risposta grezza per il blocco: {raw_output}")
        if not raw_output:
            logger.error("⚠️ Errore: Nessun testo valido restituito dall'API per l'analisi del blocco.")
            return None
        return raw_output
    except Exception as e:
        logger.error(f"⚠️ Errore nell'analisi del blocco: {e}")
        return None

def filtra_blocchi_avanzata(blocchi):
    """
    Filtra i blocchi di testo per individuare quelli critici.
    Il filtro utilizza due approcci:
      1. Controllo tramite regex (con i pattern definiti).
      2. Analisi contestuale tramite API, usando un prompt strutturato.
    Se almeno uno dei due approcci segnala il blocco come critico, esso viene restituito.
    """
    blocchi_filtrati = {}
    for i, blocco in enumerate(blocchi):
        # Verifica tramite regex
        regex_match = any(pattern.search(blocco) for pattern in compiled_patterns)
        # Analisi contestuale
        analysis = ai_analyze_block("", blocco, "")
        classification = "Non critico"  # fallback di default
        if analysis:
            try:
                result = json.loads(analysis)
                classification = result.get("classificazione", "Non critico")
            except Exception as e:
                logger.error(f"Errore nel parsing dell'analisi: {e}")
        # Se almeno uno segnala criticità, includi il blocco
        if regex_match or (classification == "Critico"):
            blocchi_filtrati[f"{i}_{blocco}"] = blocco
    return blocchi_filtrati

def process_file_content(file_content, file_extension):
    if file_extension == "html":
        soup = BeautifulSoup(file_content, "html.parser")
        blocks = [tag.get_text().strip() for tag in soup.find_all(["p", "span", "div", "li", "a", "h5"]) if tag.get_text().strip()]
        return blocks, file_content
    elif file_extension == "md":
        html_content = markdown.markdown(file_content)
        soup = BeautifulSoup(html_content, "html.parser")
        blocks = [tag.get_text().strip() for tag in soup.find_all(["p", "span", "div", "li", "a", "h5"]) if tag.get_text().strip()]
        return blocks, html_content
    return [], ""

def process_doc_file(uploaded_file):
    try:
        doc = Document(uploaded_file)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        st.error(f"Errore nell'apertura del file Word: {e}")
        st.stop()

def process_pdf_file(uploaded_file):
    try:
        pdf_reader = PdfReader(uploaded_file)
        paragraphs = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                paragraphs.extend([line.strip() for line in text.split("\n") if line.strip()])
        return paragraphs
    except Exception as e:
        st.error(f"Errore nell'apertura del file PDF: {e}")
        st.stop()

def process_html_content(html_content: str, modifications: dict, highlight: bool = False) -> str:
    """
    Applica le modifiche al contenuto HTML.

    Parametri:
      html_content (str): il contenuto HTML originale.
      modifications (dict): dizionario in cui le chiavi sono i blocchi originali da sostituire e
                            i valori sono le versioni modificate (o stringa vuota per eliminazioni).
      highlight (bool): se True, evidenzia il testo modificato racchiudendolo in un tag <mark>.
    """
    for original, new_text in modifications.items():
        replacement = f"<mark>{new_text}</mark>" if highlight and new_text else new_text
        pattern = re.escape(original)
        html_content = re.sub(pattern, replacement, html_content)
    return html_content

def process_pdf_content_with_overlay(pdf_file, modifications):
    """
    Esempio di funzione che elabora il PDF originale applicando le modifiche dei blocchi.
    In una implementazione reale si potrebbe utilizzare reportlab o un altro strumento per creare un nuovo PDF.
    Per questo esempio restituiamo semplicemente il contenuto originale del PDF.
    """
    pdf_file.seek(0)
    return pdf_file.read()

# Logica principale dell'applicazione
########################################

st.title("📄 Revisione Documenti")
st.write("Carica un file (HTML, Markdown, Word o PDF) e scegli come intervenire sul testo.")

# Selezione modalità
modalita = st.radio(
    "Modalità di revisione:",
    ("Riscrittura blocchi critici", "Conversione completa in plurale", "Blocchi critici + conversione completa"),
    help="Scegli la modalità di revisione più adatta alle tue esigenze."
)

# Checkbox per conversione globale
global_conversion = st.checkbox(
    "Applicare conversione globale in plurale",
    value=False,
    help="Seleziona per convertire l'intero documento dalla prima persona singolare alla prima persona plurale dopo le revisioni dei blocchi critici."
)

uploaded_file = st.file_uploader("📂 Seleziona un file (html, md, doc, docx, pdf)", type=["html", "md", "doc", "docx", "pdf"])

if uploaded_file is not None:
    try:
        file_bytes = uploaded_file.read()
        uploaded_file.seek(0)
        file_extension = uploaded_file.name.split('.')[-1].lower()
        st.success(f"File caricato con successo: {uploaded_file.name}")
    except Exception as e:
        st.error(f"Errore durante la lettura del file: {e}")
        st.stop()

    # Elaborazione del file (solo una volta) e salvataggio in session_state
    if "file_processed" not in st.session_state:
        if file_extension in ["html", "md"]:
            file_content = file_bytes.decode("utf-8")
            blocchi, html_content = process_file_content(file_content, file_extension)
            st.session_state.blocchi = blocchi
            st.session_state.html_content = html_content
            st.session_state.blocchi_da_revisionare = filtra_blocchi_avanzata(blocchi)
        elif file_extension in ["doc", "docx"]:
            paragraphs = process_doc_file(io.BytesIO(file_bytes))
            st.session_state.paragraphs = paragraphs
            st.session_state.blocchi_da_revisionare = filtra_blocchi_avanzata(paragraphs)
        elif file_extension == "pdf":
            paragraphs = process_pdf_file(io.BytesIO(file_bytes))
            st.session_state.paragraphs = paragraphs
            st.session_state.blocchi_da_revisionare = filtra_blocchi_avanzata(paragraphs)
        st.session_state.file_processed = True

    # Modalità "Conversione completa in plurale"
    if modalita == "Conversione completa in plurale":
        if file_extension in ["html", "md"]:
            file_content = file_bytes.decode("utf-8")
            if file_extension == "html":
                soup = BeautifulSoup(file_content, "html.parser")
                body = soup.body
                original_body_text = body.get_text(separator="\n") if body else file_content
            else:
                original_body_text = file_content

            if st.button("Genera Anteprima Conversione Completa in Plurale"):
                converted_text = ai_convert_first_singular_to_plural(original_body_text)
                st.session_state.converted_text = converted_text

            if "converted_text" in st.session_state:
                st.subheader("📌 Testo Revisionato (Conversione Completa in Plurale)")
                if "<" in st.session_state.converted_text:
                    final_html = st.session_state.converted_text
                else:
                    final_html = convert_plain_text_to_minimal_html(st.session_state.converted_text)
                st.components.v1.html(final_html, height=500, scrolling=True)
                st.download_button(
                    "📥 Scarica File Revisionato",
                    data=final_html.encode("utf-8"),
                    file_name="document_revised.html",
                    mime="text/html"
                )
        elif file_extension in ["doc", "docx"]:
            paragraphs = process_doc_file(io.BytesIO(file_bytes))
            full_text = "\n".join(paragraphs)
            if st.button("Genera Anteprima Conversione Completa in Plurale"):
                converted_text = ai_convert_first_singular_to_plural(full_text)
                st.session_state.converted_text = converted_text

            if "converted_text" in st.session_state:
                st.subheader("📌 Testo Revisionato (Conversione Completa in Plurale)")
                st.write(st.session_state.converted_text)
                new_doc = Document()
                new_doc.add_paragraph(st.session_state.converted_text)
                buffer = io.BytesIO()
                new_doc.save(buffer)
                st.download_button(
                    "📥 Scarica Documento Revisionato",
                    data=buffer.getvalue(),
                    file_name="document_revised.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        elif file_extension == "pdf":
            paragraphs = process_pdf_file(io.BytesIO(file_bytes))
            full_text = "\n".join(paragraphs)
            if st.button("Genera Anteprima Conversione Completa in Plurale"):
                converted_text = ai_convert_first_singular_to_plural(full_text)
                st.session_state.converted_text = converted_text

            if "converted_text" in st.session_state:
                st.subheader("📌 PDF Revisionato (Conversione Completa in Plurale)")
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, st.session_state.converted_text)
                buffer = io.BytesIO()
                pdf.output(buffer, 'F')
                st.download_button(
                    "📥 Scarica PDF Revisionato",
                    data=buffer.getvalue(),
                    file_name="document_revised.pdf",
                    mime="application/pdf"
                )
    # Modalità "Riscrittura blocchi critici" (o combinata)
    else:
        if st.session_state.blocchi_da_revisionare:
            with st.form("blocchi_form"):
                st.subheader("📌 Blocchi da revisionare")
                scelte_utente = {}
                if file_extension in ["html", "md"]:
                    blocchi = st.session_state.blocchi
                else:
                    blocchi = st.session_state.paragraphs

                for uid, blocco in st.session_state.blocchi_da_revisionare.items():
                    st.markdown(f"**{blocco}**")
                    azione = st.radio("Azione per questo blocco:", ["Riscrivi", "Elimina", "Ignora"], key=f"action_{uid}")
                    tono = None
                    if azione == "Riscrivi":
                        tono = st.selectbox("Scegli il tono:", list(TONE_OPTIONS.keys()), key=f"tone_{uid}")
                    scelte_utente[blocco] = {"azione": azione, "tono": tono}
                submitted = st.form_submit_button("✍️ Genera Documento Revisionato")
            if submitted:
                modifications = {}
                if file_extension in ["html", "md"]:
                    html_content = st.session_state.html_content
                    for blocco, info in scelte_utente.items():
                        if info["azione"] == "Riscrivi":
                            prev_blocco, next_blocco = extract_context(blocchi, blocco)
                            mod_blocco = ai_rewrite_text(blocco, prev_blocco, next_blocco, info["tono"])
                            modifications[blocco] = mod_blocco
                        elif info["azione"] == "Elimina":
                            modifications[blocco] = ""
                        else:
                            modifications[blocco] = blocco
                    final_content = process_html_content(html_content, modifications, highlight=True)
                    if global_conversion:
                        final_content = ai_convert_first_singular_to_plural(final_content)
                    st.success("✅ Revisione completata!")
                    st.subheader("🌍 Anteprima con Testo Revisionato")
                    st.components.v1.html(final_content, height=500, scrolling=True)
                    st.download_button(
                        "📥 Scarica HTML Revisionato",
                        data=final_content.encode("utf-8"),
                        file_name="document_revised.html",
                        mime="text/html"
                    )
                elif file_extension in ["doc", "docx"]:
                    for paragrafo, info in scelte_utente.items():
                        if info["azione"] == "Riscrivi":
                            prev_par, next_par = extract_context(blocchi, paragrafo)
                            mod_par = ai_rewrite_text(paragrafo, prev_par, next_par, info["tono"])
                            modifications[paragrafo] = mod_par
                        elif info["azione"] == "Elimina":
                            modifications[paragrafo] = ""
                        else:
                            modifications[paragrafo] = paragrafo
                    full_text = "\n".join([modifications.get(p, p) for p in blocchi])
                    if global_conversion:
                        full_text = ai_convert_first_singular_to_plural(full_text)
                    new_doc = Document()
                    new_doc.add_paragraph(full_text)
                    buffer = io.BytesIO()
                    new_doc.save(buffer)
                    st.success("✅ Revisione completata!")
                    st.subheader("🌍 Anteprima Testo (Word)")
                    st.download_button(
                        "📥 Scarica Documento Word Revisionato",
                        data=buffer.getvalue(),
                        file_name="document_revised.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                elif file_extension == "pdf":
                    for blocco, info in scelte_utente.items():
                        if info["azione"] == "Riscrivi":
                            prev_blocco, next_blocco = extract_context(blocchi, blocco)
                            mod_blocco = ai_rewrite_text(blocco, prev_blocco, next_blocco, info["tono"])
                            modifications[blocco] = mod_blocco
                        elif info["azione"] == "Elimina":
                            modifications[blocco] = ""
                        else:
                            modifications[blocco] = blocco
                    if global_conversion:
                        for key in modifications:
                            modifications[key] = ai_convert_first_singular_to_plural(modifications[key])
                    with st.spinner("🔄 Riscrittura in corso..."):
                        revised_pdf = process_pdf_content_with_overlay(io.BytesIO(file_bytes), modifications)
                    st.success("✅ Revisione completata!")
                    st.download_button(
                        "📥 Scarica PDF Revisionato",
                        data=revised_pdf,
                        file_name="document_revised.pdf",
                        mime="application/pdf"
                    )
        else:
            st.info("Non sono state trovate corrispondenze per i criteri di ricerca nel documento.")
